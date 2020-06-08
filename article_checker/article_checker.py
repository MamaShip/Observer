# -*- coding: utf-8 -*-

import requests
from threading import Lock, Thread
from time import sleep
import os
from bs4 import BeautifulSoup
from io import BytesIO
from docx import Document
from docx.shared import Cm, Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from queue import Queue
from .update_reason import *
import logging

#先声明一个 Logger 对象
logger = logging.getLogger("article_checker")
logger.setLevel(level=logging.DEBUG)
#然后指定其对应的 Handler 为 FileHandler 对象
handler = logging.FileHandler('article_checker.log')
#然后 Handler 对象单独指定了 Formatter 对象单独配置输出格式
formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
handler.setFormatter(formatter)
logger.addHandler(handler)

def default_callback(article_id, valid, backup_path=None, optionals={}):
    return

def IsValidUrl(url):
    # to do
    return True

# 一个消息队列用在Observer和Checker之间
# 消息队列的元素是 (文章id, url, 是否下载)
class Checker_Queue:
    _instance = None
    _lock = Lock()

    def __new__(cls, *args, **kwargs):
        if not cls._instance:
            with cls._lock:
                if not cls._instance:
                    cls._instance = super().__new__(cls)
        return cls._instance

    def __init__(self, max_size):
        self.q = Queue(maxsize=max_size)

    def put(self, article_id, url, download, block=True, timeout=None):
        if IsValidUrl(url):
            try:
                self.q.put((article_id, url, download), block=block, timeout=timeout)
            except:
                self.DoPutError()

    def get(self, block=True, timeout=1):
        try:
            (article_id, url, download) = self.q.get(block=block, timeout=timeout)
        except:
            self.DoGetError()
            return None, None, None
        else:
            return article_id, url, download

    def get_queue_size(self):
        return self.q.qsize()

    def DoGetError(self):
        # to do
        logger.error("checker queue get error")
        return

    def DoPutError(self):
        # to do
        logger.error("checker queue put error")
        return

# 从消息队列取元素进行判断
# 若被删除,执行删除的代码(暂未实装)
# 若存在且要下载，则下载到docx文件中
class Article_Checker(Thread):
    _instance = None
    _lock = Lock()

    def __new__(cls, *args, **kwargs):
        if not cls._instance:
            with cls._lock:
                if not cls._instance:
                    cls._instance = super().__new__(cls)
        return cls._instance

    def __init__(self, checker_queue, saving_path='', sleeping_time=1, call_back_func=default_callback):
        super(Article_Checker, self).__init__()
        self.queue = checker_queue
        self.saving_path = saving_path # 这里的saving_path是相对地址
        self.sleeping_time = sleeping_time
        self.call_back_func = call_back_func
        return

    def DoSavingFailed(self, article_id):
        # to do
        print('article {} fail to save'.format(article_id))
        self.call_back_func(article_id=article_id, valid=True, backup_path=None)
        return

    def DoSavingSucceed(self, article_id, title, saving_path): # saving_path 是 绝对地址
        # to do
        print('article {} saved'.format(article_id))
        self.call_back_func(article_id=article_id, valid=True, backup_path=saving_path, optionals={"title":title})
        return

    def DoArticleExist(self, article_id):
        # to do
        print('article {} exist'.format(article_id))
        self.call_back_func(article_id=article_id, valid=True, backup_path=None)
        return

    def DoArticleDeleted(self, article_id, url, delete_reason):
        print('article id {} deleted for {}'.format(article_id, delete_reason))
        # to do
        self.call_back_func(article_id=article_id, valid=False, 
                            backup_path=None, optionals={"reason":REASON_INACCESSIBLE})
        return

    def DoArticleDeletedWithoutDownload(self, article_id, url, delete_reason):
        # to do
        print('article id {} deleted for {}'.format(article_id, delete_reason))
        self.call_back_func(article_id=article_id, valid=False, 
                            backup_path=None, optionals={"reason":REASON_INACCESSIBLE})
        return

    def DoConnectionError(self, url):
        # to do
        print('Connection Error : {}'.format(url))
        logger.error('Connection Error : {}'.format(url))
        return

    def DoInvalidUrl(self, article_id, url):
        # to do
        print('Invalid Url : {}'.format(url))
        self.call_back_func(article_id=article_id, valid=False, 
                            backup_path=None, optionals={"reason":REASON_INVALID_URL})
        return

    def DoRequestError(self, url):
        # to do
        print('request failed for {}'.format(url))
        logger.error('Connection Error : {}'.format(url))
        return

    def run(self):
        while True:
            (article_id, url, download) = self.queue.get(block=True, timeout=1)
            #print('article {} get'.format(article_id))
            if not url:
                sleep(self.sleeping_time)
                continue
            print('article {} get'.format(article_id))
            try:
                page_soup = GetPageSoup(url, features='lxml')
            except requests.exceptions.ConnectionError:
                self.DoConnectionError(url)
                page_soup = None
            except requests.exceptions.InvalidURL:
                self.DoInvalidUrl(url, article_id)
                page_soup = None

            if not page_soup:
                sleep(self.sleeping_time)
                continue

            try:
                delete_flag, delete_reason = IsDeleted(page_soup)
            except:
                self.DoRequestError(url)
                sleep(self.sleeping_time)
                continue

            if not delete_flag:
                if not download:
                    self.DoArticleExist(article_id)
                else: # need download
                    file_name = str(article_id) + '.docx'
                    try:
                        title = Save2Doc(page_soup, self.saving_path + file_name)
                    except:
                        self.DoSavingFailed(article_id)
                        title = ''
                    else:
                        self.DoSavingSucceed(article_id, title, os.path.join(os.getcwd(), self.saving_path+file_name)) # 这里用的绝对地址
                    # Save2Doc(page_soup, self.saving_path + file_name)
                    # self.DoSavingSucceed(article_id, os.path.join(os.getcwd(), self.saving_path + file_name))
                continue

            else:
                if download:
                    self.DoArticleDeletedWithoutDownload(article_id, url, delete_reason)
                else:
                    self.DoArticleDeleted(article_id, url, delete_reason)

            sleep(self.sleeping_time)

# 读取页面内容
def GetPageContent(url, encoding='utf-8'):
    response = requests.get(url)
    content = response.content.decode(encoding=encoding, errors='ignore')
    return content

# 通过bs4 分析 xml
def GetPageSoup(url, features='lxml'):
    content = GetPageContent(url)
    if content is None:
        return None
    return BeautifulSoup(markup=content, features=features)

# 被删除的文章，会有<h3 class="weui-msg__title">XXX</h3>
# 以及红色感叹号 <i class="weui-icon-warn weui-icon_msg"></i>
# 正常文章只有<h2>标签
def IsDeleted(page_soup):
    h3_title = page_soup.find(name='h3', attrs={"class": "weui-msg__title"})
    red_icon = page_soup.find(name='i', attrs={"class": "weui-icon-warn weui-icon_msg"})
    grey_icon = page_soup.find(name='i', attrs={"class": "weui-icon-warn-gray weui-icon_msg"})
    warn_title = page_soup.find(name='div', attrs={"class": "weui-msg__title warn"})
    if red_icon and h3_title:
        return True, h3_title.text
    elif red_icon and warn_title:
        return True, warn_title.text
    elif grey_icon and warn_title:
        return True, warn_title.text
    else:
        return False, None

def SaveTextTag2Paragraph(doc, text_tag):
    cur_para = doc.add_paragraph()
    try:
        text = text_tag.text.strip()
    except:
        cur_para.add_run('获取时出错')
    else:
        cur_para.add_run(text)
    return

def InitDocStyle(doc, abc_font='Times New Roman', chn_font=u'宋体', indent_cm=0.74):
    style = doc.styles['Normal']
    style.font.name = abc_font
    style.element.rPr.rFonts.set(qn('w:eastAsia'), chn_font)
    paragraph_format = style.paragraph_format
    paragraph_format.first_line_indent = Cm(indent_cm)

# 保存图片和文字到docx文件中, 并且将文章的标题返回
def Save2Doc(page_soup, save_path, image_size=4.0):
    doc = Document()
    # init style for whole document
    InitDocStyle(doc, abc_font='Times New Roman', chn_font=u'宋体', indent_cm=0.74)
    # get title
    try:
        title = page_soup.find(name='h2').text.strip()
    except:
        title = doc.add_heading('获取文章标题时出错')
    else:
        title = doc.add_heading(title)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # get author information
    cur_para = doc.add_paragraph()
    cur_para.add_run('作者介绍：')
    author = page_soup.find(name='div', attrs={"class": "profile_inner"})
    SaveTextTag2Paragraph(doc, author)
    # get texts and pictures
    cur_para = doc.add_paragraph()
    cur_para.add_run('正文:')
    article_content_soup = page_soup.find(name='div', attrs={"id": "js_content"})
    for tag in article_content_soup.findAll(name=['p', 'section', 'img']):
        # 文字会在<p></p>之间出现
        # 还有可能在<section></section>之间出现
        if tag.name == 'p' and tag.text:
            SaveTextTag2Paragraph(doc, tag)
        if tag.name == 'section' and tag.text:
            SaveTextTag2Paragraph(doc, tag)
        # 图片会在<img></img>中出现
        if tag.name == 'img' and tag.has_attr('data-src'):
            image_io = DownloadImage(tag['data-src'])
            cur_para = doc.add_paragraph()
            if image_io is None:
                cur_para.add_run('获取图片时出错')
            else:
                cur_run = cur_para.add_run()
                cur_run.add_picture(image_io, width=Inches(image_size))
    doc.save(save_path)
    return title

# 以字节流的形式存入内存，然后再存入doc
def DownloadImage(url):
    # TO DO: 改成存字节流
    try:
        image_data = requests.get(url).content
    except:
        return None
    else:
        image_io = BytesIO()
        image_io.write(image_data)
        image_io.seek(0)
        return image_io

# 拿来测试
class Test_Class(Thread):

    def __init__(self, q):
        super(Test_Class, self).__init__()
        self.urls = [r'https://mp.weixin.qq.com/s/_dX8-fpPzpsxaKpi6Hjw6w',
                     r'https://mp.weixin.qq.com/s/Hr2XfjimJH2PJtYldTD_QA',
                     r'https://mp.weixin.qq.com/s?__biz=MzUyNDQyNTI1OQ==&mid=2247485113&idx=1&sn=fe519905e349eddd69e773769dcd5437&chksm=fa2cc7fdcd5b4eebc2ba2d9b0fe32a465b7603d93e05c1cb24edc44f1cdcf3290c03833de278&mpshare=1&srcid=0513MT9x68x6doNnABCcYwk2&sharer_sharetime=1589305468758&sharer_shareid=afd15624e89a727c2d7ee3f76ef31e5c&from=singlemessage&&sub&clicktime=1589326214&enterid=1589326214&forceh5=1&a&devicetype=android-29&version=27000e37&nettype=WIFI&abtest_cookie=AAACAA%3D%3D&lang=zh_CN&exportkey=A1AWfVCDKp%2FcNDipvHFRRlk%3D&pass_ticket=KJhWTmIcJaAEto1dcH6rJvecoQ7f6uO4KKUCYiKTukH3SEjgH%2B%2BN5CDveDdcGT8V&wx_header=1&scene=1&subscene=10000&clicktime=1589525209&enterid=1589525209',
                     r'https://mp.weixin.qq.com/s?__biz=MzU3Mjk1OTQ0Ng==&mid=2247484924&idx=1&sn=7d611b8c0a51e179cab51fd308e0a56c&chksm=fcc9ba45cbbe33535d0308c48d8d18d3ba6b18b8e0a2fbe636b3e9f0efaba6038942c98f6e70&mpshare=1&scene=2&srcid=&sharer_sharetime=1582717197787&sharer_shareid=246cb2c7250512fd9647a394d25bd429&from=timeline&key=4e4f4f0e2204deb082c29c40f853d0ea72f1deb6f474bd9c3830cb3efd14b0668fb557c9b25eabae3f65c091b1d76c4537fbddea4f24ebfa0aa067e8daadb1edd10d8b4ee5de2f5e5c278789d6ce108b&ascene=1&uin=ODg5OTA0Nzgw&devicetype=Windows+10+x64&version=62090070&lang=zh_CN&exportkey=A4bF5u75U5NZSImF8sEK6jw%3D&pass_ticket=YoLxbUZxJS4%2Fbmw6eOsgUHyiu9TwY%2BI0uEvVW6TCGknknWACHcEdBVsPGs%2FMy68Z',
                     r'https://mp.weixin.qq.com/s/q-WkvTApjcwqtgk9LY7Q-A',
                     r'https://mp.weixin.qq.com/s/aIWmD5E2y-Yg7oMltd2i8w']
        # self.urls = [r'https://mp.weixin.qq.com/s/q-WkvTApjcwqtgk9LY7Q-A']
        self.q = q

    def run(self):
        i = 0
        while len(self.urls) > 0:
            self.q.put(article_id=i, url=self.urls[0], download=True, block=True, timeout=1)
            print('article {} put'.format(i))
            self.urls.pop(0)
            sleep(3)
            i += 1

if __name__ == '__main__':
    ac = Article_Checker(Checker_Queue(max_size=100), sleeping_time=3, saving_path='')
    t = Test_Class(Checker_Queue(max_size=100))
    ac.start()
    t.start()
    t.join()
    ac.join()

